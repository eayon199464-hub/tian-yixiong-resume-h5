from pathlib import Path

from docx import Document


resume_path = Path(__file__).resolve().parents[1] / "public" / "田一雄_高级数字化产品经理_中文简历.docx"
document = Document(resume_path)

title_paragraph = next(
    (paragraph for paragraph in document.paragraphs if paragraph.text.strip() in {"田一雄", "田一雄 / Eayon"}),
    None,
)
if title_paragraph is None or len(title_paragraph.runs) != 1:
    raise RuntimeError("Expected a single-run resume name paragraph containing 田一雄")

title_paragraph.runs[0].text = "田一雄 / Eayon"

header_paragraph = document.sections[0].header.paragraphs[0]
if not header_paragraph.runs or not header_paragraph.text.startswith("田一雄"):
    raise RuntimeError("Expected the resume header to begin with 田一雄")
header_paragraph.runs[0].text = "田一雄 / Eayon"

document.save(resume_path)

print(resume_path)
